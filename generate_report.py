from docx import Document
from docx.shared import Inches
import pandas as pd
import os

# Initialize the report document
report = Document()
report.add_heading("Alpha Generation Project Report", level=0)

# Add model summary
report.add_heading("Model Summary", level=1)
model_summary = """
- Model Type: LSTM (Long Short-Term Memory)
- Layers:
  - LSTM Layer 1: 50 units (return_sequences=True)
  - LSTM Layer 2: 50 units (return_sequences=False)
  - Dense Output Layer: 1 unit (final prediction)
- Optimizer: Adam
- Loss Function: Mean Squared Error
"""
report.add_paragraph(model_summary)

# Add feature importance visualization
report.add_heading("Feature Importance Visualization", level=1)
feature_importance_image_path = "results/feature_importance.png"
if os.path.exists(feature_importance_image_path):
    report.add_picture(feature_importance_image_path, width=Inches(6))
else:
    report.add_paragraph("Feature importance visualization not found.")

# Add feature importance table
report.add_heading("Top Features Based on SHAP Values", level=1)
feature_importance_csv_path = "results/feature_importance.csv"
if os.path.exists(feature_importance_csv_path):
    feature_importance_df = pd.read_csv(feature_importance_csv_path)
    top_features = feature_importance_df.head(10)

    # Add table to the report
    table = report.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Importance'

    for _, row in top_features.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Feature'])
        row_cells[1].text = f"{row['Importance']:.6f}"
else:
    report.add_paragraph("Feature importance CSV not found.")

# Add prediction summary
report.add_heading("Prediction Example", level=1)
report.add_paragraph("The model was tested with a sequence of 60 time steps for AAPL stock prices. A sample prediction is shown below:")
predicted_price = 174.01  # Example from previous result
report.add_paragraph(f"Predicted Price: ${predicted_price:.2f}")

# Add conclusions
report.add_heading("Conclusion", level=1)
conclusion_text = """
This report presents the results of the LSTM-based model for stock price prediction. 
The SHAP values highlight the importance of recent time steps in influencing the model's predictions. 

Next steps:
- Improve data preprocessing for better feature engineering.
- Experiment with other architectures like GRU or Transformer-based models.
- Deploy the model using a scalable cloud solution for real-time predictions.
"""
report.add_paragraph(conclusion_text)

# Save the report
output_path = "results/Alpha_Generation_Project_Report.docx"
report.save(output_path)

print(f"✅ Report successfully generated and saved at {output_path}")
