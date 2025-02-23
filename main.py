"""
Robust Alpha Generation Pipeline with Comprehensive Error Handling
"""

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import shap
import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches
from sklearn.preprocessing import MinMaxScaler
from tensorflow.keras.models import Sequential, load_model
from tensorflow.keras.layers import LSTM, Dense, Input, Dropout
from tensorflow.keras.optimizers import Adam
import keras_tuner as kt
from flask import Flask, request, jsonify

# Configure logging
logging.basicConfig(filename='pipeline.log', level=logging.INFO)

CONFIG = {
    "data_path": "data/engineered_stock_data.csv",
    "model_path": "model/lstm_model.h5",
    "sequence_length": 60,
    "test_size": 0.2,
    "shap_samples": 200,
    "initial_capital": 100000,
    "risk_free_rate": 0.01,
    "enable_tuning": True,
    "tuning_trials": 10,
    "required_columns": ['Date', 'Close_AAPL', 'High_AAPL', 'Low_AAPL', 'Open_AAPL', 'Volume_AAPL', 'sentiment']
}

def main():
    """Main execution pipeline with error containment"""
    try:
        # Ensure directory structure exists
        os.makedirs("data", exist_ok=True)
        os.makedirs("model", exist_ok=True)
        os.makedirs("results", exist_ok=True)

        print("\n=== DATA PREPARATION ===")
        data, scaler = load_and_prepare_data()
        
        print("\n=== MODEL TRAINING ===")
        model = train_model(data, scaler)
        
        print("\n=== MODEL EVALUATION ===")
        actual_prices, predicted_prices = evaluate_model(model, data, scaler)
        
        print("\n=== FEATURE ANALYSIS ===")
        analyze_feature_importance(model, data, scaler)
        
        print("\n=== RISK ANALYSIS ===")
        risk_metrics = calculate_risk_metrics(actual_prices)
        
        print("\n=== TRADING SIMULATION ===")
        trading_results = simulate_trading_strategy(actual_prices, predicted_prices)
        
        print("\n=== REPORT GENERATION ===")
        generate_final_report(actual_prices, predicted_prices, risk_metrics, trading_results)
        
        print("\n=== PIPELINE COMPLETED SUCCESSFULLY ===")

    except Exception as e:
        logging.error(f"Pipeline failure: {str(e)}")
        print(f"\n❌ Critical pipeline error: {str(e)}")
        raise

def load_and_prepare_data():
    """Robust data loading with validation"""
    try:
        if not os.path.exists(CONFIG["data_path"]):
            raise FileNotFoundError(
                f"Data file not found at {CONFIG['data_path']}\n"
                f"Please ensure:\n"
                f"1. File exists in the 'data' directory\n"
                f"2. Contains required columns: {CONFIG['required_columns']}"
            )

        data = pd.read_csv(CONFIG["data_path"], parse_dates=['Date'])
        
        # Validate dataset structure
        missing_cols = set(CONFIG["required_columns"]) - set(data.columns)
        if missing_cols:
            raise ValueError(f"Missing required columns: {missing_cols}")
            
        data.set_index('Date', inplace=True)
        
        # Handle missing data
        if data.isnull().sum().any():
            data = data.ffill().bfill()
            logging.warning("Missing values filled using forward/backward fill")
            
        return data, MinMaxScaler(feature_range=(0, 1)).fit(data[['Close_AAPL']])
        
    except Exception as e:
        logging.error(f"Data preparation failed: {str(e)}")
        raise

def create_sequences(data, scaler):
    """Sequence creation with bounds checking"""
    scaled_data = scaler.transform(data[['Close_AAPL']])
    
    if len(scaled_data) < CONFIG["sequence_length"]:
        raise ValueError(
            f"Insufficient data points ({len(scaled_data)}) "
            f"for sequence length {CONFIG['sequence_length']}"
        )
    
    X, y = [], []
    for i in range(CONFIG["sequence_length"], len(scaled_data)):
        seq = scaled_data[i-CONFIG["sequence_length"]:i]
        target = scaled_data[i, 0]
        X.append(seq)
        y.append(target)
    
    X, y = np.array(X), np.array(y)
    
    if len(X) == 0:
        raise ValueError("No sequences created - check data dimensions")
    
    split = int(len(X) * (1 - CONFIG["test_size"]))
    return X[:split], X[split:], y[:split], y[split:]

def train_model(data, scaler):
    """Model training with safeguards"""
    try:
        X_train, X_test, y_train, y_test = create_sequences(data, scaler)
        
        if CONFIG["enable_tuning"]:
            model = tune_hyperparameters(X_train, y_train)
        else:
            model = build_default_model()
        
        # Validate training data
        if len(X_train) == 0 or len(y_train) == 0:
            raise ValueError("Empty training data - check sequence parameters")
            
        model.fit(X_train, y_train,
                  epochs=50,
                  batch_size=32,
                  validation_data=(X_test, y_test),
                  verbose=1)
        
        model.save(CONFIG["model_path"])
        return model
        
    except Exception as e:
        logging.error(f"Model training failed: {str(e)}")
        raise

def tune_hyperparameters(X_train, y_train):
    """Hyperparameter tuning with validation"""
    tuner = kt.RandomSearch(
        lambda hp: build_tunable_model(hp),
        objective='val_loss',
        max_trials=CONFIG["tuning_trials"],
        directory='tuning_results',
        project_name='stock_prediction'
    )
    
    tuner.search(X_train, y_train,
                 epochs=10,
                 validation_split=0.1,
                 batch_size=32,
                 verbose=0)
    
    return tuner.get_best_models(num_models=1)[0]

def build_tunable_model(hp):
    """Tunable model architecture"""
    model = Sequential([
        LSTM(hp.Int('units', 50, 300, step=50),
             return_sequences=True,
             input_shape=(CONFIG["sequence_length"], 1)),
        Dropout(hp.Float('dropout1', 0.1, 0.5, step=0.1)),
        LSTM(hp.Int('units2', 50, 200, step=50)),
        Dropout(hp.Float('dropout2', 0.1, 0.5, step=0.1)),
        Dense(1)
    ])
    model.compile(
        optimizer=Adam(hp.Float('lr', 1e-4, 1e-2, sampling='log')),
        loss='mse'
    )
    return model

def build_default_model():
    """Default model architecture"""
    return Sequential([
        LSTM(50, return_sequences=True, input_shape=(CONFIG["sequence_length"], 1)),
        Dropout(0.2),
        LSTM(50),
        Dropout(0.2),
        Dense(1)
    ])

def evaluate_model(model, data, scaler):
    """Model evaluation with safety checks"""
    try:
        _, X_test, _, y_test = create_sequences(data, scaler)
        
        if len(X_test) == 0:
            raise ValueError("No test sequences available for evaluation")
            
        predictions = model.predict(X_test)
        
        # Inverse transformations
        predictions = scaler.inverse_transform(predictions.reshape(-1, 1))
        actual_prices = scaler.inverse_transform(y_test.reshape(-1, 1))
        
        # Visualization
        plt.figure(figsize=(14, 5))
        plt.plot(actual_prices, label='Actual Prices')
        plt.plot(predictions, label='Predicted Prices')
        plt.title('Price Prediction Performance')
        plt.legend()
        plt.savefig("results/predictions.png")
        plt.close()
        
        return actual_prices.flatten(), predictions.flatten()
        
    except Exception as e:
        logging.error(f"Evaluation failed: {str(e)}")
        raise

def analyze_feature_importance(model, data, scaler):
    """SHAP analysis with proper dimension handling"""
    try:
        background = create_background_samples(data, scaler)
        
        if background.size == 0:
            raise ValueError("No background samples available for SHAP analysis")
            
        explainer = shap.GradientExplainer(
            (model.input, model.output[:, 0]),
            background
        )
        
        # Calculate SHAP values properly
        shap_values = explainer.shap_values(background)
        
        # Handle multi-output and reshape correctly
        if isinstance(shap_values, list):
            shap_values = shap_values[0]
            
        # Reshape to 2D array (samples, sequence_length)
        shap_values_2d = np.array(shap_values).squeeze()
        background_2d = background.squeeze()
        
        # Create time step features
        feature_names = [f't-{i}' for i in range(CONFIG["sequence_length"])]
        
        # Plot with correct dimensions
        plt.figure(figsize=(12, 8))
        shap.summary_plot(
            shap_values_2d,
            background_2d,
            feature_names=feature_names,
            plot_type="bar",
            show=False
        )
        plt.title("Feature Importance (SHAP Values)")
        plt.savefig("results/feature_importance.png", bbox_inches='tight')
        plt.close()
        
    except Exception as e:
        logging.error(f"SHAP analysis failed: {str(e)}")
        raise

def create_background_samples(data, scaler):
    """Safe background sample creation"""
    scaled_data = scaler.transform(data[['Close_AAPL']])
    max_samples = len(scaled_data) - CONFIG["sequence_length"]
    sample_size = min(CONFIG["shap_samples"], max_samples)
    
    if sample_size <= 0:
        return np.array([])
        
    indices = np.random.choice(max_samples, sample_size, replace=False)
    return np.array([scaled_data[i:i+CONFIG["sequence_length"]] for i in indices])

def calculate_risk_metrics(prices):
    """Risk calculation with numerical safeguards"""
    try:
        if len(prices) < 2:
            return {'error': 'Insufficient price data for risk calculation'}
            
        with np.errstate(divide='ignore', invalid='ignore'):
            returns = np.diff(prices) / prices[:-1]
            returns = np.nan_to_num(returns)
            
        return {
            'var_95': np.percentile(returns, 5) if len(returns) > 0 else 0,
            'cvar_95': returns[returns <= np.percentile(returns, 5)].mean() if len(returns) > 0 else 0,
            'max_drawdown': (prices / np.maximum.accumulate(prices) - 1).min()
        }
        
    except Exception as e:
        logging.error(f"Risk calculation error: {str(e)}")
        return {'error': str(e)}

def simulate_trading_strategy(actual_prices, predicted_prices):
    """Trading simulation with bounds checking"""
    try:
        min_length = min(len(actual_prices), len(predicted_prices))
        actual_prices = actual_prices[:min_length]
        predicted_prices = predicted_prices[:min_length]
        
        signals = np.where(predicted_prices[1:] > predicted_prices[:-1], 1, -1)
        positions = np.zeros_like(actual_prices)
        positions[1:] = signals
        
        portfolio = CONFIG["initial_capital"] * np.cumprod(1 + positions[:-1] * (actual_prices[1:]/actual_prices[:-1] - 1))
        
        return {
            'final_value': portfolio[-1] if len(portfolio) > 0 else CONFIG["initial_capital"],
            'returns': portfolio / CONFIG["initial_capital"] - 1
        }
        
    except Exception as e:
        logging.error(f"Trading simulation failed: {str(e)}")
        return {'error': str(e)}

def generate_final_report(actual, predicted, risks, trading):
    """Report generation with error handling"""
    try:
        doc = Document()
        doc.add_heading('Alpha Generation Report', 0)
        
        # Prediction Plot
        doc.add_heading('Price Predictions', level=1)
        doc.add_picture('results/predictions.png', width=Inches(6))
        
        # Feature Importance
        doc.add_heading('Feature Importance', level=1)
        doc.add_picture('results/feature_importance.png', width=Inches(6))
        
        # Risk Metrics
        doc.add_heading('Risk Metrics', level=1)
        risk_table = doc.add_table(rows=1, cols=3)
        risk_table.style = 'LightShading'
        row = risk_table.rows[0].cells
        row[0].text = f"95% VaR: {risks.get('var_95', 0):.4f}"
        row[1].text = f"95% CVaR: {risks.get('cvar_95', 0):.4f}"
        row[2].text = f"Max Drawdown: {risks.get('max_drawdown', 0):.4f}"
        
        # Trading Results
        doc.add_heading('Trading Performance', level=1)
        doc.add_paragraph(f"Initial Capital: ${CONFIG['initial_capital']:,.2f}")
        doc.add_paragraph(f"Final Value: ${trading.get('final_value', 0):,.2f}")
        
        doc.save('results/final_report.docx')
        
    except Exception as e:
        logging.error(f"Report generation failed: {str(e)}")
        raise

if __name__ == "__main__":
    main()